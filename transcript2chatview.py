#!/usr/bin/env python3
"""
Microsoft Teams DOCX文字起こしをChatView形式のマークダウンに変換するツール

使い方:
    python transcript2chatview.py input.docx -o output.md
    python transcript2chatview.py input.docx --merge-speaker   # 同一話者の連続発言を結合
    python transcript2chatview.py input.docx --no-timestamp    # タイムスタンプ非表示
    python transcript2chatview.py input.docx --no-icon         # アイコン絵文字非表示
    python transcript2chatview.py input.docx --merge-speaker --no-timestamp --no-icon  # 複数オプション併用
"""

import argparse
import re
import base64
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def extract_paragraph_images(docx_file, output_dir=None, use_files=True):
    """
    DOCXファイルから段落ごとに画像を抽出
    
    Args:
        docx_file: DOCXファイルのパス
        output_dir: 画像ファイルを保存するディレクトリ（use_files=Trueの場合）
        use_files: Trueの場合はファイルとして保存、Falseの場合はBase64エンコード
        
    Returns:
        dict: {paragraph_index: {'path': str} or {'data_uri': str, 'content_type': str}}
    """
    doc = Document(docx_file)
    
    # すべての画像リレーションシップを取得
    image_parts = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype.lower():
            image_parts[rel.rId] = rel.target_part
    
    # 画像保存用ディレクトリを作成
    if use_files and output_dir:
        icons_dir = Path(output_dir) / 'icons'
        icons_dir.mkdir(parents=True, exist_ok=True)
    
    # 段落ごとに画像を抽出
    paragraph_images = {}
    
    for para_idx, para in enumerate(doc.paragraphs):
        # 段落内のdrawing要素を検索
        para_element = para._element
        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        # drawing要素を検索
        drawings = para_element.findall('.//w:drawing', ns)
        
        for drawing in drawings:
            # blip要素から画像IDを取得
            blips = drawing.findall('.//a:blip', ns)
            
            for blip in blips:
                embed_id = blip.get(qn('r:embed'))
                if embed_id and embed_id in image_parts:
                    image_part = image_parts[embed_id]
                    image_data = image_part.blob
                    content_type = image_part.content_type
                    
                    if use_files and output_dir:
                        # ファイルとして保存
                        ext = content_type.split('/')[-1]
                        icon_filename = f"speaker_{para_idx:03d}.{ext}"
                        icon_path = icons_dir / icon_filename
                        
                        with open(icon_path, 'wb') as f:
                            f.write(image_data)
                        
                        paragraph_images[para_idx] = {
                            'path': f"icons/{icon_filename}"
                        }
                    else:
                        # Base64エンコード
                        base64_image = base64.b64encode(
                            image_data).decode('utf-8')
                        data_uri = f"data:{content_type};base64,"
                        data_uri += f"{base64_image}"
                        
                        paragraph_images[para_idx] = {
                            'data_uri': data_uri,
                            'content_type': content_type
                        }
                    break  # 最初の画像のみ使用
            
            if para_idx in paragraph_images:
                break  # 最初のdrawingのみ使用
    
    return paragraph_images


def parse_teams_docx_simple(docx_file, output_dir=None, use_icon_files=True):
    """
    Teams通常形式のDOCXファイルをパース
    話者名 タイムスタンプ
    本文
    の形式に対応（1つの段落内に改行で含まれる場合も対応）
    画像アイコンも抽出して話者と紐づけ
    
    Args:
        docx_file: DOCXファイルパス
        output_dir: アイコン画像を保存するディレクトリ
        use_icon_files: Trueの場合は画像ファイルとして保存、
                        Falseの場合はBase64埋め込み
    """
    doc = Document(docx_file)
    
    # 段落ごとの画像を抽出
    paragraph_images = extract_paragraph_images(
        docx_file, output_dir, use_files=use_icon_files)
    
    transcript = []
    speaker_icons = {}  # 話者名 -> path or data_uri のマッピング
    
    # 話者情報のパターン
    speaker_pattern = re.compile(r'^(.+?)\s{2,}(\d+:\d+)')
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 段落内の改行で分割
        lines = text.split('\n')
        
        # 最初の行が話者情報かチェック
        if lines and speaker_pattern.match(lines[0]):
            first_line = lines[0]
            speaker_match = speaker_pattern.match(first_line)
            
            if speaker_match:
                speaker = speaker_match.group(1).strip()
                timestamp = '00:' + speaker_match.group(2)  # 00:を追加
                
                # この段落に画像があれば、話者と紐づけ
                if para_idx in paragraph_images:
                    # 初めて見る話者の場合のみアイコンを登録
                    if speaker not in speaker_icons:
                        img_info = paragraph_images[para_idx]
                        if 'path' in img_info:
                            speaker_icons[speaker] = img_info['path']
                        else:
                            speaker_icons[speaker] = img_info['data_uri']
                
                # 残りの行を本文として結合
                content_lines = lines[1:]
                content = '\n'.join(content_lines).strip()
                
                if content:  # 本文がある場合のみ追加
                    # 話者に紐づいたアイコンを使用
                    icon_ref = speaker_icons.get(speaker, '')
                    
                    transcript.append({
                        'start': timestamp + '.000',
                        'end': timestamp + '.000',
                        'speaker': speaker,
                        'icon': icon_ref,
                        'text': content
                    })
    
    return transcript


def parse_webvtt_from_docx(docx_file):
    """
    Teams DOCXファイル（WEBVTT形式を含む）をパースして構造化データに変換
    
    Args:
        docx_file: DOCXファイルのパス
        
    Returns:
        list: [{'start': str, 'end': str, 'speaker': str, 'text': str}, ...]
    """
    doc = Document(docx_file)
    transcript = []
    
    # すべての段落を結合してテキストとして取得
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # デバッグ: 最初の500文字を表示
    # print(f"DEBUG: Full text preview:\n{full_text[:500]}\n")
    # print(f"DEBUG: Total lines: {len(full_text.split(chr(10)))}\n")
    
    # WEBVTT形式のパターン: <v 話者名>テキスト</v>
    # タイムスタンプ行とテキスト行を抽出
    lines = full_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # タイムスタンプ行を検出
        timestamp_match = re.match(
            r'(\d+:\d+:\d+\.\d+)\s*-->\s*(\d+:\d+:\d+\.\d+)', line)
        
        if timestamp_match:
            start_time = timestamp_match.group(1)
            end_time = timestamp_match.group(2)
            
            # 次の行がVTT形式のテキスト
            i += 1
            if i < len(lines):
                text_line = lines[i].strip()
                
                # <v 話者名>テキスト</v> の形式をパース
                vtt_match = re.match(r'<v\s+([^>]+)>(.*?)</v>', text_line)
                
                if vtt_match:
                    speaker = vtt_match.group(1).strip()
                    text = vtt_match.group(2).strip()
                    
                    transcript.append({
                        'start': start_time,
                        'end': end_time,
                        'speaker': speaker,
                        'text': text
                    })
        
        i += 1
    
    return transcript


def parse_teams_docx(docx_file):
    """
    Teams DOCXファイルをパースして構造化データに変換
    
    Args:
        docx_file: DOCXファイルのパス
        
    Returns:
        list: [{'start': str, 'end': str, 'speaker': str, 'text': str}, ...]
    """
    # まず通常のTeams形式を試す
    transcript = parse_teams_docx_simple(docx_file)
    if transcript:
        return transcript
    
    # 次にWEBVTT形式を試す
    transcript = parse_webvtt_from_docx(docx_file)
    if transcript:
        return transcript
    
    # WEBVTT形式でない場合、従来の形式でパース
    doc = Document(docx_file)
    transcript = []
    current_entry = {}
    state = 'waiting_timestamp'
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 空行をスキップ
        if not text:
            continue
        
        # タイムスタンプ行を検出
        timestamp_match = re.match(r'(\d+:\d+:\d+\.\d+)\s*-->\s*(\d+:\d+:\d+\.\d+)', text)
        
        if timestamp_match:
            # 前のエントリを保存
            if current_entry and current_entry.get('text'):
                transcript.append(current_entry)
            
            # 新しいエントリを開始
            current_entry = {
                'start': timestamp_match.group(1),
                'end': timestamp_match.group(2),
                'speaker': None,
                'text': ''
            }
            state = 'waiting_speaker'
            
        elif state == 'waiting_speaker' and current_entry.get('speaker') is None:
            # 話者名行
            current_entry['speaker'] = text
            state = 'waiting_text'
            
        elif state == 'waiting_text' and current_entry.get('speaker'):
            # テキスト内容行（複数行の可能性あり）
            if current_entry['text']:
                current_entry['text'] += ' ' + text
            else:
                current_entry['text'] = text
    
    # 最後のエントリを追加
    if current_entry and current_entry.get('text'):
        transcript.append(current_entry)
    
    return transcript


def merge_consecutive_speakers(transcript):
    """
    同一話者の連続した発言を結合
    
    Args:
        transcript: パースされた文字起こしデータ
        
    Returns:
        list: 結合後のデータ
    """
    if not transcript:
        return []
    
    merged = []
    current = transcript[0].copy()
    
    for entry in transcript[1:]:
        if entry['speaker'] == current['speaker']:
            # 同じ話者なら結合
            current['text'] += ' ' + entry['text']
            current['end'] = entry['end']  # 終了時刻を更新
        else:
            # 違う話者なら保存して新規開始
            merged.append(current)
            current = entry.copy()
    
    # 最後のエントリを追加
    merged.append(current)
    
    return merged


def get_speaker_icon(speaker_name, speaker_index):
    """
    話者に応じたアイコン絵文字を返す
    
    Args:
        speaker_name: 話者名
        speaker_index: 話者の出現順（0始まり）
        
    Returns:
        str: 絵文字アイコン
    """
    # 話者ごとに異なるアイコンを割り当て
    icons = ['👨', '👩', '🧑', '👴', '👵', '👦', '👧', '🧔', '👱', '👨‍💼']
    return icons[speaker_index % len(icons)]


def convert_to_chatview_markdown(transcript, show_timestamp=True, show_icon=True):
    """
    パースした文字起こしをChatView形式のマークダウンに変換
    
    Args:
        transcript: パースされたデータ
        show_timestamp: タイムスタンプを表示するか
        show_icon: アイコンを表示するか
        
    Returns:
        str: ChatView形式のマークダウン
    """
    markdown_lines = []
    
    # 話者ごとにuserとassistantを交互に割り当て
    speaker_roles = {}
    speaker_icons = {}
    role_toggle = ['ai', 'me']
    role_index = 0
    
    for entry in transcript:
        speaker = entry['speaker']
        text = entry['text'].strip()
        timestamp = entry['start']
        
        # entryにアイコンが含まれているかチェック
        entry_icon = entry.get('icon', '')
        
        # 初出の話者にロールとアイコンを割り当て
        if speaker not in speaker_roles:
            speaker_roles[speaker] = role_toggle[role_index % 2]
            # entryにアイコンがあればそれを使用、なければデフォルト絵文字
            if entry_icon:
                # ファイルパスかBase64かを判定
                if entry_icon.startswith('data:image'):
                    # Base64画像の場合はHTMLのimg形式で埋め込む
                    img_tag = f'<img src="{entry_icon}" '
                    img_tag += 'width="20" height="20" />'
                    speaker_icons[speaker] = img_tag
                else:
                    # ファイルパスの場合もimg形式
                    img_tag = f'<img src="{entry_icon}" '
                    img_tag += 'width="20" height="20" />'
                    speaker_icons[speaker] = img_tag
            else:
                speaker_icons[speaker] = get_speaker_icon(
                    speaker, role_index)
            role_index += 1
        
        role = speaker_roles[speaker]
        icon = speaker_icons[speaker]
        
        # ChatView形式で出力
        if show_icon and icon:
            header = f'@{role}[{icon} {speaker}]'
        else:
            header = f'@{role}[{speaker}]'
        
        if show_timestamp:
            header += f'{{{timestamp}}}'
        
        markdown_lines.append(header)
        
        markdown_lines.append(text)
        markdown_lines.append('')
    
    return '\n'.join(markdown_lines)


def main():
    parser = argparse.ArgumentParser(
        description='Microsoft Teams DOCX文字起こしをChatView形式に変換'
    )
    parser.add_argument(
        'input',
        type=Path,
        help='入力DOCXファイル'
    )
    parser.add_argument(
        '-o', '--output',
        type=Path,
        help='出力マークダウンファイル（省略時は標準出力）'
    )
    parser.add_argument(
        '--merge-speaker',
        action='store_true',
        help='同一話者の連続発言を結合'
    )
    parser.add_argument(
        '--no-timestamp',
        action='store_true',
        help='タイムスタンプを非表示'
    )
    parser.add_argument(
        '--no-icon',
        action='store_true',
        help='アイコン絵文字を非表示'
    )
    parser.add_argument(
        '--icon-files',
        action='store_true',
        help='アイコンを別ファイルとして保存（Base64埋め込みを避ける）'
    )
    
    args = parser.parse_args()
    
    # ファイル存在チェック
    if not args.input.exists():
        print(f'エラー: ファイルが見つかりません: {args.input}')
        return 1
    
    # 出力ディレクトリを決定
    if args.output:
        output_dir = args.output.parent
    else:
        output_dir = args.input.parent
    
    # DOCXファイルをパース
    print(f'文字起こしファイルを読み込んでいます: {args.input}')
    transcript = parse_teams_docx_simple(
        args.input,
        output_dir=output_dir,
        use_icon_files=args.icon_files
    )
    print(f'  → {len(transcript)}件のエントリを検出')
    
    # オプション: 連続話者を結合
    if args.merge_speaker:
        print('同一話者の連続発言を結合しています...')
        transcript = merge_consecutive_speakers(transcript)
        print(f'  → {len(transcript)}件に結合')
    
    # ChatView形式に変換
    print('ChatView形式のマークダウンに変換しています...')
    markdown = convert_to_chatview_markdown(
        transcript,
        show_timestamp=not args.no_timestamp,
        show_icon=not args.no_icon
    )
    
    # 出力
    if args.output:
        args.output.write_text(markdown, encoding='utf-8')
        print(f'変換完了: {args.output}')
    else:
        print('\n--- 変換結果 ---\n')
        print(markdown)
    
    return 0


if __name__ == '__main__':
    exit(main())

