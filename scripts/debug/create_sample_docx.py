#!/usr/bin/env python3
"""
テスト用のTeams文字起こしDOCXファイルを作成
"""

from docx import Document

def create_sample_transcript():
    doc = Document()
    
    # サンプルデータ
    entries = [
        ("0:0:5.100 --> 0:0:8.250", "田中太郎", "おはようございます。今日の議題は新プロジェクトについてです。"),
        ("0:0:8.500 --> 0:0:10.300", "田中太郎", "まず予算から確認していきましょう。"),
        ("0:0:10.800 --> 0:0:14.500", "鈴木花子", "はい、資料を共有しますね。"),
        ("0:0:15.200 --> 0:0:18.900", "鈴木花子", "今年度の予算は前年比120%で計画しています。"),
        ("0:0:19.300 --> 0:0:22.100", "佐藤次郎", "質問なのですが、この予算には人件費も含まれていますか？"),
        ("0:0:22.500 --> 0:0:25.800", "鈴木花子", "はい、含まれています。詳細は3ページ目をご覧ください。"),
        ("0:0:26.200 --> 0:0:28.500", "田中太郎", "ありがとうございます。では次の議題に移ります。"),
        ("0:0:29.000 --> 0:0:32.400", "田中太郎", "スケジュールについて確認したいと思います。"),
    ]
    
    for timestamp, speaker, text in entries:
        # タイムスタンプ
        doc.add_paragraph(timestamp)
        # 話者名
        doc.add_paragraph(speaker)
        # 発言内容
        doc.add_paragraph(text)
        # 空行
        doc.add_paragraph()
    
    doc.save('sample_transcript.docx')
    print("✓ sample_transcript.docx を作成しました")

if __name__ == '__main__':
    create_sample_transcript()

