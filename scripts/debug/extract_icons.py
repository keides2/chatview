#!/usr/bin/env python3
"""
DOCXファイルから画像（アイコン）を抽出するスクリプト
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def extract_images_from_docx(docx_file, output_dir=None):
    """
    DOCXファイルから画像を抽出
    
    Args:
        docx_file: DOCXファイルのパス
        output_dir: 画像の保存先ディレクトリ（Noneの場合は表示のみ）
    """
    doc = Document(docx_file)
    
    if output_dir:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
    
    # ドキュメント内のすべての画像リレーションシップを取得
    image_parts = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype.lower():
            image_parts[rel.rId] = rel.target_part
            # print(f"Found image: {rel.rId} -> {rel.target_ref}")
    
    print(f"\nTotal images found: {len(image_parts)}")
    
    # 段落と画像の関連付けを抽出
    paragraph_images = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        # 段落内のdrawing要素を検索
        para_element = para._element
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
              'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
              'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
        
        # drawing要素を検索
        drawings = para_element.findall('.//w:drawing', ns)
        
        for drawing in drawings:
            # blip要素から画像IDを取得
            blips = drawing.findall('.//a:blip', ns)
            
            for blip in blips:
                embed_id = blip.get(qn('r:embed'))
                if embed_id and embed_id in image_parts:
                    image_part = image_parts[embed_id]
                    
                    # 段落のテキストを取得
                    para_text = para.text[:100] if para.text else "(empty)"
                    
                    info = {
                        'paragraph_index': para_idx,
                        'paragraph_text': para_text,
                        'image_id': embed_id,
                        'content_type': image_part.content_type,
                        'image_data': image_part.blob
                    }
                    
                    paragraph_images.append(info)
                    
                    print(f"\nParagraph {para_idx}: {para_text}")
                    print(f"  Image ID: {embed_id}")
                    print(f"  Content Type: {image_part.content_type}")
                    print(f"  Size: {len(image_part.blob)} bytes")
                    
                    # 画像を保存
                    if output_dir:
                        ext = image_part.content_type.split('/')[-1]
                        filename = f"icon_{para_idx:03d}_{embed_id}.{ext}"
                        filepath = output_path / filename
                        with open(filepath, 'wb') as f:
                            f.write(image_part.blob)
                        print(f"  Saved to: {filepath}")
    
    return paragraph_images


def analyze_speaker_icons(docx_file):
    """
    話者とアイコンの関連付けを分析
    """
    print(f"\n{'='*60}")
    print(f"Analyzing: {docx_file}")
    print(f"{'='*60}\n")
    
    images = extract_images_from_docx(docx_file, output_dir="extracted_icons")
    
    if not images:
        print("\n⚠️ No images found in the document")
        return
    
    print(f"\n{'='*60}")
    print(f"Summary: Found {len(images)} image(s) in the document")
    print(f"{'='*60}")
    
    # 話者名のパターンを抽出
    doc = Document(docx_file)
    speaker_pattern = re.compile(r'^(.+?)\s{2,}(\d+:\d+)')
    
    speakers_with_icons = {}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 段落内の改行で分割
        lines = text.split('\n')
        
        # 最初の行が話者情報かチェック
        if lines and speaker_pattern.match(lines[0]):
            first_line = lines[0]
            speaker_match = speaker_pattern.match(first_line)
            
            if speaker_match:
                speaker = speaker_match.group(1).strip()
                
                # この段落にアイコンがあるかチェック
                for img in images:
                    if img['paragraph_index'] == i:
                        if speaker not in speakers_with_icons:
                            speakers_with_icons[speaker] = []
                        speakers_with_icons[speaker].append(img)
    
    if speakers_with_icons:
        print("\n📊 Speakers with icons:")
        for speaker, icons in speakers_with_icons.items():
            print(f"\n  👤 {speaker}: {len(icons)} icon(s)")


if __name__ == '__main__':
    import re
    
    if len(sys.argv) < 2:
        print("Usage: python extract_icons.py <docx_file>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if not Path(docx_file).exists():
        print(f"Error: File not found: {docx_file}")
        sys.exit(1)
    
    analyze_speaker_icons(docx_file)
