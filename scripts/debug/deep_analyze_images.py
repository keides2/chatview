#!/usr/bin/env python3
"""
DOCXファイルの画像を詳細に分析するスクリプト
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re


def deep_analyze_images(docx_file):
    """
    DOCXファイル内の画像を徹底的に分析
    """
    print(f"\n{'='*60}")
    print(f"Deep Image Analysis: {docx_file}")
    print(f"{'='*60}\n")
    
    doc = Document(docx_file)
    
    # 1. Document.partからすべてのリレーションシップを確認
    print("1. Document Part Relationships:")
    print("-" * 60)
    for rel_id, rel in doc.part.rels.items():
        print(f"  ID: {rel_id}")
        print(f"  Type: {rel.reltype}")
        print(f"  Target: {rel.target_ref}")
        if "image" in rel.target_ref.lower():
            print(f"  ✅ IMAGE FOUND!")
            print(f"  Content-Type: {rel.target_part.content_type}")
            print(f"  Size: {len(rel.target_part.blob)} bytes")
        print()
    
    # 2. 各段落内の画像を詳細分析
    print("\n2. Paragraph-level Image Analysis:")
    print("-" * 60)
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text[:80] if para.text else "(empty)"
        
        # 段落のXML要素を直接確認
        para_element = para._element
        
        # すべてのdrawing要素を検索
        drawings = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        
        if drawings:
            print(f"\n[Paragraph {para_idx}]")
            print(f"  Text: {para_text}")
            print(f"  Drawings found: {len(drawings)}")
            
            for draw_idx, drawing in enumerate(drawings):
                print(f"\n  Drawing {draw_idx}:")
                
                # drawing内のblip要素を検索（画像参照）
                blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                
                for blip in blips:
                    # 埋め込み画像のID
                    embed_id = blip.get(qn('r:embed'))
                    # リンク画像のID
                    link_id = blip.get(qn('r:link'))
                    
                    if embed_id:
                        print(f"    Embedded Image ID: {embed_id}")
                        # リレーションシップから画像情報を取得
                        if embed_id in doc.part.rels:
                            img_part = doc.part.rels[embed_id].target_part
                            print(f"    Content-Type: {img_part.content_type}")
                            print(f"    Size: {len(img_part.blob)} bytes")
                    
                    if link_id:
                        print(f"    Linked Image ID: {link_id}")
                
                # extent（サイズ情報）
                extents = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                for extent in extents:
                    cx = extent.get('cx')
                    cy = extent.get('cy')
                    if cx and cy:
                        # EMUからピクセルに変換（914400 EMU = 1 inch, 96 DPI）
                        width_px = int(cx) / 914400 * 96
                        height_px = int(cy) / 914400 * 96
                        print(f"    Size: {width_px:.1f}x{height_px:.1f} px")
    
    # 3. run内の画像を確認
    print("\n\n3. Run-level Image Analysis:")
    print("-" * 60)
    
    image_count = 0
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            run_element = run._element
            
            # run内のdrawing要素
            drawings = run_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            
            if drawings:
                para_text = para.text[:50] if para.text else "(empty)"
                print(f"\n[Para {para_idx}, Run {run_idx}]")
                print(f"  Text: {para_text}...")
                print(f"  Run text: '{run.text}'")
                
                for drawing in drawings:
                    blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed_id = blip.get(qn('r:embed'))
                        if embed_id and embed_id in doc.part.rels:
                            image_count += 1
                            img_part = doc.part.rels[embed_id].target_part
                            print(f"  ✅ Image {image_count}: {embed_id}")
                            print(f"     Type: {img_part.content_type}")
                            print(f"     Size: {len(img_part.blob)} bytes")
    
    print(f"\n{'='*60}")
    print(f"Total images found: {image_count}")
    print(f"{'='*60}\n")
    
    return image_count


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python deep_analyze_images.py <docx_file>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if not Path(docx_file).exists():
        print(f"Error: File not found: {docx_file}")
        sys.exit(1)
    
    deep_analyze_images(docx_file)
