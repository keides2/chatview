#!/usr/bin/env python3
"""DOCXファイルの内容をデバッグ表示"""
from docx import Document
import sys

if len(sys.argv) < 2:
    print("使い方: python debug_docx.py <docxファイル>")
    sys.exit(1)

docx_file = sys.argv[1]
doc = Document(docx_file)

print(f"ファイル: {docx_file}")
print(f"総段落数: {len(doc.paragraphs)}")
print("\n=== 最初の30段落 ===\n")

for i, para in enumerate(doc.paragraphs[:30]):
    text = para.text.strip()
    if text:  # 空行でない場合のみ表示
        print(f"[{i:3d}] {repr(text)}")
    else:
        print(f"[{i:3d}] (空行)")
