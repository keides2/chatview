#!/usr/bin/env python3
"""
DOCXファイルの構造を詳細に分析するスクリプト
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import re


def analyze_docx_structure(docx_file):
    """
    DOCXファイルの構造を詳細に分析
    """
    print(f"\n{'='*60}")
    print(f"Analyzing: {docx_file}")
    print(f"{'='*60}\n")
    
    doc = Document(docx_file)
    
    # 1. 画像の確認
    print("📷 Image Analysis:")
    print("-" * 60)
    image_parts = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_parts[rel.rId] = rel.target_part
            print(f"  Image ID: {rel.rId}")
            print(f"  Target: {rel.target_ref}")
            print(f"  Content Type: {rel.target_part.content_type}")
            print(f"  Size: {len(rel.target_part.blob)} bytes")
            print()
    
    if not image_parts:
        print("  ⚠️ No images found\n")
    else:
        print(f"  ✅ Total: {len(image_parts)} image(s)\n")
    
    # 2. 段落の詳細分析
    print("\n📝 Paragraph Analysis:")
    print("-" * 60)
    
    speaker_pattern = re.compile(r'^(.+?)\s{2,}(\d+:\d+)')
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 最初の100文字を表示
        preview = text[:100] + ('...' if len(text) > 100 else '')
        print(f"\n[Paragraph {idx}]")
        print(f"  Text: {preview}")
        
        # 段落内のruns（テキストの断片）を分析
        print(f"  Runs: {len(para.runs)}")
        for run_idx, run in enumerate(para.runs):
            if run.text.strip():
                run_preview = run.text[:50] + ('...' if len(run.text) > 50 else '')
                print(f"    Run {run_idx}: '{run_preview}'")
                
                # runに画像が含まれているか確認
                run_element = run._element
                drawings = run_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                if drawings:
                    print(f"      🖼️ Contains drawing element")
                
                # inline画像を確認
                blips = run_element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                for blip in blips:
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        print(f"      🖼️ Image embedded: {embed}")
        
        # 話者パターンにマッチするか確認
        lines = text.split('\n')
        if lines and speaker_pattern.match(lines[0]):
            match = speaker_pattern.match(lines[0])
            speaker = match.group(1).strip()
            timestamp = match.group(2)
            print(f"  👤 Speaker detected: {speaker}")
            print(f"  ⏰ Timestamp: {timestamp}")
    
    # 3. スタイル情報
    print("\n\n🎨 Styles:")
    print("-" * 60)
    styles_used = set()
    for para in doc.paragraphs:
        if para.style and para.style.name:
            styles_used.add(para.style.name)
    
    for style in sorted(styles_used):
        print(f"  - {style}")
    
    print(f"\n{'='*60}\n")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python analyze_docx.py <docx_file>")
        print("\nExample:")
        print("  python analyze_docx.py sample.docx")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if not Path(docx_file).exists():
        print(f"❌ Error: File not found: {docx_file}")
        sys.exit(1)
    
    analyze_docx_structure(docx_file)
