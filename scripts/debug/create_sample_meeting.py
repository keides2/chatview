#!/usr/bin/env python3
"""
テスト用のサンプルDOCXファイルを作成
Teams会議の文字起こし形式（アイコン付き）
"""

from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path


def create_sample_docx_with_icons():
    """
    アイコン画像付きのサンプルDOCXファイルを作成
    """
    doc = Document()
    
    # サンプルの会議内容
    conversations = [
        {
            'speaker': 'TANAKA Taro 田中 太郎',
            'timestamp': '02:27',
            'text': 'おはようございます。本日の議題について説明します。'
        },
        {
            'speaker': 'SUZUKI Hanako 鈴木 花子',
            'timestamp': '02:45',
            'text': 'ありがとうございます。質問があります。'
        },
        {
            'speaker': 'TANAKA Taro 田中 太郎',
            'timestamp': '03:10',
            'text': 'はい、どうぞ。'
        },
    ]
    
    for conv in conversations:
        # 話者名とタイムスタンプの段落
        p = doc.add_paragraph()
        
        # アイコン用のプレースホルダー（絵文字で代用）
        # 実際のTeams DOCXでは画像が埋め込まれますが、
        # テストでは絵文字を使用
        speaker_name = conv['speaker']
        if 'TANAKA' in speaker_name:
            icon = '👨'
        elif 'SUZUKI' in speaker_name:
            icon = '👩'
        else:
            icon = '👤'
        
        # 話者情報を追加
        run = p.add_run(f"{icon} {speaker_name}  {conv['timestamp']}\n")
        run.font.size = Pt(11)
        
        # 発言内容を追加
        run = p.add_run(conv['text'])
        run.font.size = Pt(10)
        
        # 空行を追加
        doc.add_paragraph()
    
    # ファイルを保存
    output_path = Path('scripts/debug/sample_meeting.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ Sample DOCX created: {output_path}")
    print(f"   Contains {len(conversations)} conversation entries")
    
    return output_path


if __name__ == '__main__':
    create_sample_docx_with_icons()
